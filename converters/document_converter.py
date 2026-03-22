"""Document conversion module."""

from __future__ import annotations

import logging
import os
import re
from typing import Any

logger = logging.getLogger(__name__)


class ConversionError(Exception):
    """Raised when document conversion fails."""


class ValidationError(Exception):
    """Raised when input validation fails."""


# ---------------------------------------------------------------------------
# Font registration (DejaVu Sans — full Cyrillic support)
# ---------------------------------------------------------------------------

_fonts_registered = False


def _fonts_dir() -> str:
    return os.path.join(os.path.dirname(os.path.dirname(__file__)), 'static', 'fonts')


def _register_fonts() -> None:
    global _fonts_registered
    if _fonts_registered:
        return

    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.pdfmetrics import registerFontFamily
    from reportlab.pdfbase.ttfonts import TTFont

    d = _fonts_dir()
    pdfmetrics.registerFont(TTFont('DejaVu', os.path.join(d, 'DejaVuSans.ttf')))
    pdfmetrics.registerFont(TTFont('DejaVu-Bold', os.path.join(d, 'DejaVuSans-Bold.ttf')))
    pdfmetrics.registerFont(TTFont('DejaVu-Italic', os.path.join(d, 'DejaVuSans-Oblique.ttf')))
    pdfmetrics.registerFont(TTFont('DejaVu-BoldItalic', os.path.join(d, 'DejaVuSans-BoldOblique.ttf')))
    registerFontFamily(
        'DejaVu',
        normal='DejaVu', bold='DejaVu-Bold',
        italic='DejaVu-Italic', boldItalic='DejaVu-BoldItalic',
    )
    _fonts_registered = True


# ---------------------------------------------------------------------------
# reportlab helpers for PDF generation
# ---------------------------------------------------------------------------

def _build_pdf(output_path: str, story_items: list) -> None:
    """Build a PDF document from a list of reportlab Flowables."""
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate

    _register_fonts()
    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        leftMargin=50, rightMargin=50, topMargin=50, bottomMargin=50,
    )
    doc.build(story_items)


def _get_styles():
    """Return a dict of ParagraphStyles with DejaVu font."""
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT

    _register_fonts()
    base = dict(fontName='DejaVu', leading=18, alignment=TA_LEFT)
    return {
        'body':  ParagraphStyle('body', fontSize=11, **base),
        'h1':    ParagraphStyle('h1', fontSize=22, fontName='DejaVu-Bold', leading=28, spaceAfter=10, alignment=TA_LEFT),
        'h2':    ParagraphStyle('h2', fontSize=18, fontName='DejaVu-Bold', leading=24, spaceAfter=8, alignment=TA_LEFT),
        'h3':    ParagraphStyle('h3', fontSize=14, fontName='DejaVu-Bold', leading=20, spaceAfter=6, alignment=TA_LEFT),
        'code':  ParagraphStyle('code', fontSize=9, fontName='DejaVu', leading=12, backColor='#f4f4f4', alignment=TA_LEFT),
    }


def _escape(text: str) -> str:
    """Escape HTML special characters for reportlab Paragraph."""
    return (
        text
        .replace('&', '&amp;')
        .replace('<', '&lt;')
        .replace('>', '&gt;')
        .replace('"', '&quot;')
    )


# ---------------------------------------------------------------------------
# Converters
# ---------------------------------------------------------------------------

def convert_pdf_to_docx(input_path: str, output_path: str) -> dict[str, Any]:
    """Convert PDF to DOCX using pdf2docx."""
    from pdf2docx import Converter

    original_size = os.path.getsize(input_path)
    try:
        cv = Converter(input_path)
        cv.convert(output_path)
        cv.close()
    except Exception as exc:
        raise ConversionError(f"PDF → DOCX: {exc}") from exc

    converted_size = os.path.getsize(output_path)
    logger.info("PDF -> DOCX: %d -> %d bytes", original_size, converted_size)
    return {'original_size': original_size, 'converted_size': converted_size}


def convert_docx_to_pdf(input_path: str, output_path: str) -> dict[str, Any]:
    """Convert DOCX to PDF using python-docx + reportlab."""
    from docx import Document
    from reportlab.platypus import Paragraph, Spacer

    original_size = os.path.getsize(input_path)
    styles = _get_styles()

    try:
        doc = Document(input_path)
        story = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                story.append(Spacer(1, 6))
                continue

            style_name = para.style.name.lower() if para.style else ''
            if 'heading 1' in style_name:
                story.append(Paragraph(_escape(text), styles['h1']))
            elif 'heading 2' in style_name:
                story.append(Paragraph(_escape(text), styles['h2']))
            elif 'heading 3' in style_name:
                story.append(Paragraph(_escape(text), styles['h3']))
            else:
                # Preserve bold/italic from runs
                parts = []
                for run in para.runs:
                    t = _escape(run.text)
                    if not t:
                        continue
                    if run.bold and run.italic:
                        t = f'<b><i>{t}</i></b>'
                    elif run.bold:
                        t = f'<b>{t}</b>'
                    elif run.italic:
                        t = f'<i>{t}</i>'
                    parts.append(t)
                story.append(Paragraph(''.join(parts) or _escape(text), styles['body']))

        if not story:
            story.append(Paragraph('(пустой документ)', styles['body']))

        _build_pdf(output_path, story)

    except ConversionError:
        raise
    except Exception as exc:
        raise ConversionError(f"DOCX → PDF: {exc}") from exc

    converted_size = os.path.getsize(output_path)
    logger.info("DOCX -> PDF: %d -> %d bytes", original_size, converted_size)
    return {'original_size': original_size, 'converted_size': converted_size}


def convert_pdf_to_images(
    input_path: str,
    output_dir: str,
    prefix: str,
    fmt: str = 'PNG',
    dpi: int = 200,
    max_pages: int = 50,
) -> dict[str, Any]:
    """Convert PDF pages to images using PyMuPDF."""
    import fitz

    original_size = os.path.getsize(input_path)
    ext = '.png' if fmt.upper() == 'PNG' else '.jpg'

    try:
        doc = fitz.open(input_path)
        page_count = min(len(doc), max_pages)
        filenames = []
        zoom = dpi / 72.0
        matrix = fitz.Matrix(zoom, zoom)

        for i in range(page_count):
            pix = doc[i].get_pixmap(matrix=matrix)
            fname = f"{prefix}_page_{i + 1}{ext}"
            pix.save(os.path.join(output_dir, fname))
            filenames.append(fname)

        doc.close()
    except Exception as exc:
        raise ConversionError(f"PDF → изображения: {exc}") from exc

    logger.info("PDF -> %d images (%s, %d DPI)", page_count, fmt, dpi)
    return {'original_size': original_size, 'page_count': page_count, 'filenames': filenames}


def convert_images_to_pdf(input_paths: list[str], output_path: str) -> dict[str, Any]:
    """Merge multiple images into a single PDF using Pillow."""
    from PIL import Image

    if not input_paths:
        raise ValidationError("Не указаны изображения")

    total_original = sum(os.path.getsize(p) for p in input_paths)

    try:
        images = []
        for path in input_paths:
            img = Image.open(path)
            if img.mode != 'RGB':
                img = img.convert('RGB')
            images.append(img)

        first = images[0]
        rest = images[1:] if len(images) > 1 else []
        first.save(output_path, 'PDF', save_all=True, append_images=rest)
        for img in images:
            img.close()

    except (ValidationError, ConversionError):
        raise
    except Exception as exc:
        raise ConversionError(f"Изображения → PDF: {exc}") from exc

    converted_size = os.path.getsize(output_path)
    logger.info("Images (%d) -> PDF: %d -> %d bytes", len(input_paths), total_original, converted_size)
    return {'original_size': total_original, 'converted_size': converted_size, 'page_count': len(input_paths)}


def convert_markdown_to_pdf(input_path: str, output_path: str) -> dict[str, Any]:
    """Convert Markdown to PDF via markdown + reportlab."""
    import markdown
    from reportlab.platypus import Paragraph, Preformatted, Spacer

    original_size = os.path.getsize(input_path)
    styles = _get_styles()

    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            md_text = f.read()

        html_body = markdown.markdown(md_text, extensions=['tables', 'fenced_code'])
        story = _html_to_story(html_body, styles)

        if not story:
            story.append(Paragraph('(пустой документ)', styles['body']))

        _build_pdf(output_path, story)

    except ConversionError:
        raise
    except Exception as exc:
        raise ConversionError(f"Markdown → PDF: {exc}") from exc

    converted_size = os.path.getsize(output_path)
    logger.info("MD -> PDF: %d -> %d bytes", original_size, converted_size)
    return {'original_size': original_size, 'converted_size': converted_size}


def convert_markdown_to_html(input_path: str, output_path: str) -> dict[str, Any]:
    """Convert Markdown to a standalone HTML document."""
    import markdown

    original_size = os.path.getsize(input_path)

    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            md_text = f.read()

        body = markdown.markdown(md_text, extensions=['tables', 'fenced_code'])
        html = _wrap_html(body)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html)

    except Exception as exc:
        raise ConversionError(f"Markdown → HTML: {exc}") from exc

    converted_size = os.path.getsize(output_path)
    logger.info("MD -> HTML: %d -> %d bytes", original_size, converted_size)
    return {'original_size': original_size, 'converted_size': converted_size}


def convert_html_to_pdf(input_path: str, output_path: str) -> dict[str, Any]:
    """Convert HTML to PDF using reportlab."""
    from reportlab.platypus import Paragraph, Spacer

    original_size = os.path.getsize(input_path)
    styles = _get_styles()

    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            html = f.read()

        # Extract body content
        body_match = re.search(r'<body[^>]*>(.*)</body>', html, re.DOTALL | re.IGNORECASE)
        body = body_match.group(1) if body_match else html

        story = _html_to_story(body, styles)
        if not story:
            story.append(Paragraph('(пустой документ)', styles['body']))

        _build_pdf(output_path, story)

    except ConversionError:
        raise
    except Exception as exc:
        raise ConversionError(f"HTML → PDF: {exc}") from exc

    converted_size = os.path.getsize(output_path)
    logger.info("HTML -> PDF: %d -> %d bytes", original_size, converted_size)
    return {'original_size': original_size, 'converted_size': converted_size}


def convert_txt_to_pdf(input_path: str, output_path: str) -> dict[str, Any]:
    """Convert plain text to PDF using reportlab."""
    from reportlab.platypus import Preformatted

    original_size = os.path.getsize(input_path)
    styles = _get_styles()

    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            text = f.read()

        story = [Preformatted(text, styles['body'])]
        _build_pdf(output_path, story)

    except Exception as exc:
        raise ConversionError(f"TXT → PDF: {exc}") from exc

    converted_size = os.path.getsize(output_path)
    logger.info("TXT -> PDF: %d -> %d bytes", original_size, converted_size)
    return {'original_size': original_size, 'converted_size': converted_size}


# ---------------------------------------------------------------------------
# OCR
# ---------------------------------------------------------------------------

def ocr_pdf_to_text(input_path: str, output_path: str, lang: str = '') -> dict[str, Any]:
    """Extract text from a scanned PDF using OCR (Tesseract).

    Converts each PDF page to an image, runs OCR, saves combined text.
    """
    import fitz

    try:
        import pytesseract
    except ImportError:
        raise ConversionError("pytesseract не установлен. Установите: pip install pytesseract")

    from PIL import Image
    import io as _io

    original_size = os.path.getsize(input_path)

    # Auto-detect available languages
    if not lang:
        try:
            available = pytesseract.get_languages()
            lang = 'rus+eng' if 'rus' in available else 'eng'
        except Exception:
            lang = 'eng'

    try:
        doc = fitz.open(input_path)
        page_count = min(len(doc), 50)
        all_text = []

        for i in range(page_count):
            pix = doc[i].get_pixmap(dpi=300)
            img = Image.open(_io.BytesIO(pix.tobytes('png')))
            text = pytesseract.image_to_string(img, lang=lang)
            all_text.append(f"--- Страница {i + 1} ---\n{text.strip()}")
            img.close()

        doc.close()

        result_text = '\n\n'.join(all_text)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(result_text)

    except ConversionError:
        raise
    except Exception as exc:
        raise ConversionError(f"OCR: {exc}") from exc

    converted_size = os.path.getsize(output_path)
    logger.info("OCR PDF (%d pages): %d -> %d bytes", page_count, original_size, converted_size)
    return {'original_size': original_size, 'converted_size': converted_size, 'page_count': page_count}


def ocr_pdf_to_docx(input_path: str, output_path: str, lang: str = '') -> dict[str, Any]:
    """Convert scanned PDF to DOCX via OCR — each page becomes a section."""
    import fitz

    try:
        import pytesseract
    except ImportError:
        raise ConversionError("pytesseract не установлен. Установите: pip install pytesseract")

    from PIL import Image
    import io as _io
    from docx import Document
    from docx.shared import Pt

    original_size = os.path.getsize(input_path)

    if not lang:
        try:
            available = pytesseract.get_languages()
            lang = 'rus+eng' if 'rus' in available else 'eng'
        except Exception:
            lang = 'eng'

    try:
        pdf = fitz.open(input_path)
        page_count = min(len(pdf), 50)

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        for i in range(page_count):
            if i > 0:
                doc.add_page_break()

            pix = pdf[i].get_pixmap(dpi=300)
            img = Image.open(_io.BytesIO(pix.tobytes('png')))
            text = pytesseract.image_to_string(img, lang=lang)
            img.close()

            doc.add_heading(f'Страница {i + 1}', level=2)
            for line in text.strip().split('\n'):
                doc.add_paragraph(line)

        pdf.close()
        doc.save(output_path)

    except ConversionError:
        raise
    except Exception as exc:
        raise ConversionError(f"OCR → DOCX: {exc}") from exc

    converted_size = os.path.getsize(output_path)
    logger.info("OCR PDF -> DOCX (%d pages): %d -> %d bytes", page_count, original_size, converted_size)
    return {'original_size': original_size, 'converted_size': converted_size, 'page_count': page_count}


# ---------------------------------------------------------------------------
# HTML helpers (for markdown/html -> reportlab story)
# ---------------------------------------------------------------------------

def _html_to_story(html_body: str, styles: dict) -> list:
    """Convert simple HTML body to a list of reportlab Flowables."""
    from reportlab.platypus import Paragraph, Preformatted, Spacer

    story = []
    # Split by block-level tags
    parts = re.split(r'(</?(?:h[1-3]|p|pre|ul|ol|li|br|hr)[^>]*>)', html_body)

    current_tag = None
    buffer = []

    for part in parts:
        part = part.strip()
        if not part:
            continue

        tag_match = re.match(r'^<(/?)(h[1-3]|p|pre|ul|ol|li|br|hr)', part, re.IGNORECASE)
        if tag_match:
            closing = tag_match.group(1)
            tag = tag_match.group(2).lower()

            if closing:
                # Closing tag — flush buffer
                text = ' '.join(buffer).strip()
                buffer = []
                if text:
                    if current_tag in ('h1', 'h2', 'h3'):
                        story.append(Paragraph(_strip_tags(text), styles[current_tag]))
                    elif current_tag == 'pre':
                        story.append(Preformatted(_strip_tags(text), styles['code']))
                    else:
                        # Allow <b>, <i>, <code>, <strong>, <em> inline tags for reportlab
                        clean = _clean_inline_html(text)
                        story.append(Paragraph(clean, styles['body']))
                current_tag = None
            else:
                if tag in ('br', 'hr'):
                    story.append(Spacer(1, 8))
                else:
                    current_tag = tag
        else:
            buffer.append(part)

    # Remaining text
    text = ' '.join(buffer).strip()
    if text:
        story.append(Paragraph(_clean_inline_html(text), styles.get(current_tag, styles['body'])))

    return story


def _strip_tags(html: str) -> str:
    """Remove all HTML tags."""
    return re.sub(r'<[^>]+>', '', html).strip()


def _clean_inline_html(html: str) -> str:
    """Keep only safe inline tags for reportlab: b, i, u, strong, em, br."""
    # Replace <strong> with <b>, <em> with <i>
    html = re.sub(r'<strong>', '<b>', html, flags=re.IGNORECASE)
    html = re.sub(r'</strong>', '</b>', html, flags=re.IGNORECASE)
    html = re.sub(r'<em>', '<i>', html, flags=re.IGNORECASE)
    html = re.sub(r'</em>', '</i>', html, flags=re.IGNORECASE)
    # Remove all other tags except b, i, u, br
    html = re.sub(r'<(?!/?(b|i|u|br)\b)[^>]+>', '', html)
    return html.strip()


def _wrap_html(body: str) -> str:
    """Wrap HTML body in a full document (for MD→HTML export only)."""
    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
    body {{ font-family: sans-serif; font-size: 14px; line-height: 1.6; margin: 40px; color: #222; }}
    h1 {{ font-size: 24px; margin: 20px 0 10px; }}
    h2 {{ font-size: 20px; margin: 18px 0 8px; }}
    h3 {{ font-size: 16px; margin: 14px 0 6px; }}
    pre {{ background: #f4f4f4; padding: 12px; font-size: 12px; }}
    code {{ background: #f4f4f4; padding: 2px 4px; font-size: 13px; }}
    table {{ border-collapse: collapse; width: 100%; margin: 12px 0; }}
    th, td {{ border: 1px solid #ccc; padding: 8px; text-align: left; }}
    th {{ background: #f0f0f0; }}
</style>
</head>
<body>
{body}
</body>
</html>"""
